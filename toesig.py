import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from collections import defaultdict
from datetime import datetime, timedelta
import uuid
import logging

# Set wide layout
st.set_page_config(layout="wide")

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Load admin periods CSV
df = pd.read_csv("admin periods.csv")

# Convert CSV to dictionary
data = {}
for day in df.columns[1:]:
    data[day] = {}
    for _, row in df.iterrows():
        period = row['Period']
        teachers = row[day]
        teacher_list = [t.strip() for t in teachers.split(',') if t.strip()]
        data[day][period] = teacher_list

# Get unique educators
all_educators = sorted(set(t for day in data for period in data[day] for t in data[day][period]))
logger.debug(f"All educators: {all_educators}")

# SMT teachers
SMT_TEACHERS = [
    "AR VISAGIE", "C MATTHYS", "G ZEALAND", "J KLEIN",
    "Y COETZEE", "I DIEDERICKS", "R BRANDT", "E CLOETE", "J SAAL",
    "ML MATTHYS", "M CLOETE", "P GELDERBLOM", "D VAN EEDEN"
]

# Excluded teachers
EXCLUDED_TEACHERS = []

# Initialize session state
if "absence_counts" not in st.session_state:
    st.session_state.absence_counts = defaultdict(int)
if "usage_counts" not in st.session_state:
    st.session_state.usage_counts = defaultdict(int)
if "usage_timestamps" not in st.session_state:
    st.session_state.usage_timestamps = defaultdict(list)
if "daily_substitutes" not in st.session_state:
    st.session_state.daily_substitutes = defaultdict(set)
if "absence_timestamps" not in st.session_state:
    st.session_state.absence_timestamps = defaultdict(list)
if "return_periods" not in st.session_state:
    st.session_state.return_periods = {}
if "leave_periods" not in st.session_state:
    st.session_state.leave_periods = {}
if "last_start_period" not in st.session_state:
    st.session_state.last_start_period = None

# CSS styling
st.markdown("""
    <style>
        .main {
            background-color: #F5F7FA;
            padding: 20px;
            border-radius: 10px;
            max-width: 95vw !important;
            width: 100%;
            margin: 0 auto;
        }
        .stButton>button {
            background-color: #005B99;
            color: white;
            border-radius: 5px;
            padding: 10px 20px;
            font-weight: bold;
        }
        .stButton>button:hover {
            background-color: #007ACC;
        }
        .generate-button>button, .download-button>button {
            background-color: #28A745 !important;
            color: white !important;
            border-radius: 5px !important;
            padding: 10px 20px !important;
            font-weight: bold !important;
            border: none !important;
        }
        .generate-button>button:hover, .download-button>button:hover {
            background-color: #218838 !important;
        }
        h1, h2, h3 {
            color: #003087;
            text-align: center;
        }
        .section {
            background-color: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            margin-bottom: 20px;
            max-width: 95vw !important;
            width: 100%;
            margin: 0 auto;
        }
        .table-container {
            overflow-x: auto;
            width: 100%;
        }
        .wide-table {
            width: 100%;
            max-width: 1600px !important;
            margin: 0 auto;
            font-size: 18px;
            border-collapse: collapse;
            word-spacing: 3px;
            white-space: nowrap;
            background-color: #FFFFFF;
        }
        .wide-table th, .wide-table td {
            padding: 15px !important;
            text-align: center;
            border: 1px solid #005B99;
        }
        .wide-table th {
            background-color: #005B99;
            color: white;
            font-weight: bold;
            font-size: 20px;
        }
        .wide-table td {
            background-color: #E6F0FA;
            color: #003087;
        }
        .wide-table tr:nth-child(even) td {
            background-color: #D1E3F5;
        }
        .wide-table tr:hover td {
            background-color: #B3D4FC;
            transition: background-color 0.2s;
        }
        .stSelectbox, .stMultiSelect, .stNumberInput {
            width: 90% !important;
            max-width: none !important;
        }
        .stColumn > div {
            display: flex;
            justify-content: center;
        }
    </style>
""", unsafe_allow_html=True)

# Streamlit app
st.title("SAUL DAMON HIGH SCHOOL TOESIGROOSTER")
st.markdown("<div class='main'>", unsafe_allow_html=True)

# Input Section
st.markdown("<div class='section'>", unsafe_allow_html=True)
st.subheader("Afwesigheid en Skedule Konfigurasie")
col1, col2, col3, col4, col5 = st.columns(5)
with col1:
    absent_educators = st.multiselect("Kies Afwesige Opvoeders", all_educators, key="absent_educators")
with col2:
    selected_day = st.selectbox("Kies Dag", list(data.keys()), key="selected_day")
with col3:
    day_layout = st.selectbox("Kies Dag Uitleg", ["2 Periodes Voor Eet/Break", "3 Periodes Voor Eet/Break"], key="day_layout")
with col4:
    start_period = st.selectbox("Begin Periode", [f"Periode {i}" for i in range(1, 8)], key="start_period")
with col5:
    num_periods = st.number_input("Aantal Periodes", min_value=1, max_value=7, value=7, key="num_periods")

# Clear inputs
if st.button("Maak Alle Insette Skoon"):
    st.session_state.clear()
    st.rerun()

# Check if start period changed and clear invalid periods
start_period_idx = int(start_period.split()[-1])
if st.session_state.last_start_period != start_period:
    period_options = ["Volle Dag Afwesig"] + [f"Periode {i}" for i in range(1, 8)]
    leave_period_options = ["Geen Vroeë Vertrek"] + [f"Periode {i}" for i in range(1, 8)]
    
    # Clear invalid return periods
    for educator in list(st.session_state.return_periods.keys()):
        if st.session_state.return_periods[educator] not in period_options:
            logger.debug(f"Clearing invalid return period for {educator}: {st.session_state.return_periods[educator]}")
            del st.session_state.return_periods[educator]
    
    # Clear invalid leave periods
    for educator in list(st.session_state.leave_periods.keys()):
        if st.session_state.leave_periods[educator] not in leave_period_options:
            logger.debug(f"Clearing invalid leave period for {educator}: {st.session_state.leave_periods[educator]}")
            del st.session_state.leave_periods[educator]
    
    st.session_state.last_start_period = start_period

# Specify return periods for absent teachers
period_options = ["Volle Dag Afwesig"] + [f"Periode {i}" for i in range(1, 8)]
if absent_educators:
    st.subheader("Spesifiseer Terugkeer Periodes vir Afwesige Opvoeders")
    for educator in absent_educators:
        key = f"return_{educator}"
        default_value = st.session_state.return_periods.get(educator, "Volle Dag Afwesig")
        # Ensure default_value is in period_options
        if default_value not in period_options:
            default_value = "Volle Dag Afwesig"
            logger.debug(f"Resetting return period for {educator} to 'Volle Dag Afwesig' as {default_value} is invalid")
        return_period = st.selectbox(
            f"Wanneer keer {educator} terug?",
            period_options,
            index=period_options.index(default_value),
            key=key
        )
        st.session_state.return_periods[educator] = return_period
        logger.debug(f"Selected return period for {educator}: {return_period}")

# Specify leave periods for non-absent teachers
non_absent_educators = [t for t in all_educators if t not in absent_educators]
if non_absent_educators:
    st.subheader("Spesifiseer Vroeë Vertrek vir Ander Opvoeders")
    leaving_educators = st.multiselect("Kies Opvoeders wat Vroeg Vertrek", non_absent_educators, key="leaving_educators")
    if leaving_educators:
        leave_period_options = ["Geen Vroeë Vertrek"] + [f"Periode {i}" for i in range(1, 8)]
        for educator in leaving_educators:
            key = f"leave_{educator}"
            default_value = st.session_state.leave_periods.get(educator, "Geen Vroeë Vertrek")
            # Ensure default_value is in leave_period_options
            if default_value not in leave_period_options:
                default_value = "Geen Vroeë Vertrek"
                logger.debug(f"Resetting leave period for {educator} to 'Geen Vroeë Vertrek' as {default_value} is invalid")
            leave_period = st.selectbox(
                f"Wanneer vertrek {educator}?",
                leave_period_options,
                index=leave_period_options.index(default_value),
                key=key
            )
            st.session_state.leave_periods[educator] = leave_period
            logger.debug(f"Selected leave period for {educator}: {leave_period}")
else:
    st.session_state.leave_periods = {}

st.markdown("</div>", unsafe_allow_html=True)

# Generate schedule
days = list(data.keys())
current_day_idx = days.index(selected_day)
next_day = days[(current_day_idx + 1) % len(days)]

teaching_periods = []
full_schedule = []
current_period_idx = start_period_idx
periods_added = 0
teaching_periods_count = 0
teaching_periods_since_pouse1 = 0

periods_before_eat = 2 if day_layout == "2 Periodes Voor Eet/Break" else 3
periods_before_pouse2 = 3 if day_layout == "2 Periodes Voor Eet/Break" else 2

while periods_added < num_periods:
    period_num = ((current_period_idx - 1) % 7) + 1 if current_period_idx > 7 else current_period_idx
    period_name = f"Periode {period_num}"
    if current_period_idx > 7:
        period_name += " (Dag 2)"
    teaching_periods.append(period_name)
    full_schedule.append(period_name)
    teaching_periods_count += 1
    teaching_periods_since_pouse1 += 1
    periods_added += 1
    current_period_idx += 1
    
    if teaching_periods_count == periods_before_eat and periods_added < num_periods:
        full_schedule.extend(["Eet", "POUSE 1"])
        teaching_periods_since_pouse1 = 0
    if teaching_periods_since_pouse1 == periods_before_pouse2 and periods_added < num_periods:
        full_schedule.append("POUSE 2")
        teaching_periods_since_pouse1 = 0

# Select substitute teacher
def select_substitute(selected_day, period, absent_educators, used_teachers, start_period_idx, daily_used_teachers):
    try:
        period_num = int(period.split()[1].split(" ")[0])
        period_position = teaching_periods.index(period) + 1
        total_periods_first_day = 7 - start_period_idx + 1
        day_to_use = selected_day if period_position <= total_periods_first_day else next_day
        scheduled_teachers = data.get(day_to_use, {}).get(f"Period {period_num}", [])
        
        available_teachers = [
            t for t in scheduled_teachers
            if t not in absent_educators and t not in EXCLUDED_TEACHERS and t not in used_teachers
        ]
        
        if not available_teachers:
            return "OPDEEL"
        
        # Prioritize teachers not used today
        unused_teachers = [t for t in available_teachers if t not in daily_used_teachers]
        non_smt_unused = [t for t in unused_teachers if t not in SMT_TEACHERS]
        smt_unused = [t for t in unused_teachers if t in SMT_TEACHERS]
        
        if non_smt_unused:
            substitute = non_smt_unused[0]
        elif smt_unused and period not in [p for p in teaching_periods if smt_unused[0] in st.session_state.daily_substitutes[p]]:
            substitute = smt_unused[0]
        else:
            non_smt_used = [t for t in available_teachers if t not in SMT_TEACHERS]
            smt_used = [t for t in available_teachers if t in SMT_TEACHERS and period not in [p for p in teaching_periods if t in st.session_state.daily_substitutes[p]]]
            substitute = non_smt_used[0] if non_smt_used else smt_used[0] if smt_used else "OPDEEL"
        
        if substitute != "OPDEEL":
            used_teachers.add(substitute)
            st.session_state.daily_substitutes[period].add(substitute)
            st.session_state.usage_counts[substitute] += 1
            st.session_state.usage_timestamps[substitute].append((datetime.now(), day_to_use))
        return substitute
    except Exception as e:
        logger.error(f"Error in select_substitute: {str(e)}")
        return "OPDEEL"

# Generate substitution schedule
st.session_state.daily_substitutes = defaultdict(set)
daily_used_teachers = set()
unique_columns = ["Afwesige Opvoeders"] + full_schedule
all_teachers = absent_educators + [t for t in st.session_state.leave_periods.keys() if st.session_state.leave_periods[t] != "Geen Vroeë Vertrek"]
num_rows = max(2, len(all_teachers) + 1)
table_data = [["" for _ in range(len(unique_columns))] for _ in range(num_rows)]
table_data[0] = unique_columns

if all_teachers:
    for row_idx, teacher in enumerate(all_teachers, 1):
        table_data[row_idx][0] = teacher
else:
    table_data[1][0] = "Geen"

period_order = [p for p in full_schedule if p in teaching_periods]
for row_idx, teacher in enumerate(all_teachers, 1):
    is_absent = teacher in absent_educators
    return_period = st.session_state.return_periods.get(teacher, "Volle Dag Afwesig") if is_absent else None
    leave_period = st.session_state.leave_periods.get(teacher, "Geen Vroeë Vertrek") if not is_absent else None
    
    return_idx = int(return_period.split()[-1]) if is_absent and return_period != "Volle Dag Afwesig" else 8
    leave_idx = int(leave_period.split()[-1]) if not is_absent and leave_period != "Geen Vroeë Vertrek" else 8
    
    for col_idx, period in enumerate(full_schedule, 1):
        if period not in teaching_periods:
            table_data[row_idx][col_idx] = ""
        else:
            # Extract the absolute period number
            period_num = int(period.split()[1].split(" ")[0])
            if is_absent:
                if return_period != "Volle Dag Afwesig" and period_num >= return_idx:
                    table_data[row_idx][col_idx] = "AANWESIG"
                else:
                    current_absent = [
                        t for t in absent_educators
                        if st.session_state.return_periods.get(t, "Volle Dag Afwesig") == "Volle Dag Afwesig" or
                        (st.session_state.return_periods[t] != "Volle Dag Afwesig" and period_num < int(st.session_state.return_periods[t].split()[-1]))
                    ]
                    substitute = select_substitute(selected_day, period, current_absent, set(), start_period_idx, daily_used_teachers)
                    table_data[row_idx][col_idx] = substitute
                    if substitute != "OPDEEL":
                        daily_used_teachers.add(substitute)
            else:
                if leave_period != "Geen Vroeë Vertrek" and period_num >= leave_idx:
                    current_absent = [
                        t for t in st.session_state.leave_periods.keys()
                        if st.session_state.leave_periods[t] != "Geen Vroeë Vertrek" and period_num >= int(st.session_state.leave_periods[t].split()[-1])
                    ] + [
                        t for t in absent_educators
                        if st.session_state.return_periods.get(t, "Volle Dag Afwesig") == "Volle Dag Afwesig" or
                        (st.session_state.return_periods[t] != "Volle Dag Afwesig" and period_num < int(st.session_state.return_periods[t].split()[-1]))
                    ]
                    substitute = select_substitute(selected_day, period, current_absent, set(), start_period_idx, daily_used_teachers)
                    table_data[row_idx][col_idx] = substitute
                    if substitute != "OPDEEL":
                        daily_used_teachers.add(substitute)
                else:
                    table_data[row_idx][col_idx] = "AANWESIG"

# Substitution Schedule Table
st.markdown("<div class='section'>", unsafe_allow_html=True)
st.subheader("TOESIGROOSTER")
df_schedule = pd.DataFrame(table_data[1:], columns=unique_columns)
df_schedule.index = df_schedule.index + 1

unique_cols = []
seen = {}
for col in unique_columns:
    if col in seen:
        seen[col] += 1
        unique_cols.append(f"{col}_{seen[col]}")
    else:
        seen[col] = 0
        unique_cols.append(col)
df_schedule.columns = unique_cols

st.markdown("<div class='table-container'><div class='wide-table'>", unsafe_allow_html=True)
st.table(df_schedule.style.set_properties(**{
    'background-color': '#E6F0FA',
    'color': '#003087',
    'border': '1px solid #005B99',
    'padding': '15px',
    'text-align': 'center',
    'font-size': '18px',
    'word-spacing': '3px',
    'white-space': 'nowrap'
}).set_table_styles([
    {'selector': 'th', 'props': [('background-color', '#005B99'), ('color', 'white'), ('font-weight', 'bold'), ('font-size', '20px'), ('padding', '15px')]},
    {'selector': 'tr:nth-child(even)', 'props': [('background-color', '#D1E3F5')]},
    {'selector': 'tr:hover', 'props': [('background-color', '#B3D4FC')]}
]))
st.markdown("</div></div>", unsafe_allow_html=True)
st.markdown("</div>", unsafe_allow_html=True)

# Available Teachers Table
st.markdown("<div class='section'>", unsafe_allow_html=True)
st.subheader("Beskikbare Opvoeders per Periode")
available_data = {}
for period in teaching_periods:
    period_num = int(period.split()[1].split(" ")[0])
    period_position = teaching_periods.index(period) + 1
    total_periods_first_day = 7 - start_period_idx + 1
    day_to_use = selected_day if period_position <= total_periods_first_day else next_day
    current_absent = [
        t for t in absent_educators
        if st.session_state.return_periods.get(t, "Volle Dag Afwesig") == "Volle Dag Afwesig" or
        (st.session_state.return_periods[t] != "Volle Dag Afwesig" and period_num < int(st.session_state.return_periods[t].split()[-1]))
    ] + [
        t for t in st.session_state.leave_periods.keys()
        if st.session_state.leave_periods[t] != "Geen Vroeë Vertrek" and period_num >= int(st.session_state.leave_periods[t].split()[-1])
    ]
    available_teachers = [
        t for t in data[day_to_use][f"Period {period_num}"]
        if t not in current_absent and t not in EXCLUDED_TEACHERS and t not in st.session_state.daily_substitutes[period]
    ]
    available_data[period] = ", ".join(available_teachers or ["Geen"])
df_available = pd.DataFrame(list(available_data.items()), columns=["Periode", "Beskikbare Opvoeders"])
df_available.index = df_available.index + 1
st.table(df_available.style.set_properties(**{
    'background-color': '#E6F0FA',
    'color': '#003087',
    'border': '1px solid #005B99',
    'padding': '10px',
    'text-align': 'center',
    'font-size': '16px'
}).set_table_styles([{
    'selector': 'th',
    'props': [('background-color', '#005B99'), ('color', 'white'), ('font-weight', 'bold'), ('font-size', '16px'), ('padding', '10px')]
}]))
st.markdown("</div>", unsafe_allow_html=True)

# Generate document
st.markdown("<div class='section'>", unsafe_allow_html=True)
st.markdown("<div class='generate-button'>", unsafe_allow_html=True)
if st.button("Genereer TOESIGROOSTER"):
    current_date = datetime.now()
    for educator in absent_educators:
        if st.session_state.return_periods.get(educator, "Volle Dag Afwesig") != f"Periode {start_period_idx}":
            st.session_state.absence_counts[educator] += 1
            st.session_state.absence_timestamps[educator].append(current_date)
    for educator in st.session_state.leave_periods:
        if st.session_state.leave_periods[educator] != "Geen Vroeë Vertrek":
            st.session_state.absence_counts[educator] += 1
            st.session_state.absence_timestamps[educator].append(current_date)
    
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    doc.add_heading("SAUL DAMON HIGH SCHOOL", 0).alignment = 1
    p = doc.add_paragraph()
    run = p.add_run(f"TOESIGROOSTER VIR {selected_day} ({datetime.now().strftime('%d/%m/%Y')})")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 48, 135)
    p.alignment = 1
    doc.add_paragraph(f"Dag Uitleg: {day_layout} | Begin by: {start_period} | Aantal Periodes: {num_periods}").alignment = 1
    
    num_rows = len(all_teachers) + 1 if all_teachers else 2
    table = doc.add_table(rows=num_rows, cols=len(unique_columns))
    table.style = 'Table Grid'
    table.autofit = False
    
    def set_table_borders(table):
        tbl = table._element
        tbl_pr = tbl.tblPr
        tbl_borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '005B99')
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)
    
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(unique_columns):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.paragraph_format.space_after = Pt(0)
        cell_fill = OxmlElement('w:shd')
        cell_fill.set(qn('w:fill'), '005B99')
        hdr_cells[i]._element.get_or_add_tcPr().append(cell_fill)
    
    page_width = section.page_width - section.left_margin - section.right_margin
    col_width = page_width / len(unique_columns)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width
    
    if all_teachers:
        for row_idx, teacher in enumerate(all_teachers, 1):
            for col_idx in range(len(unique_columns)):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = table_data[row_idx][col_idx]
                if row_idx % 2 == 0:
                    cell_fill = OxmlElement('w:shd')
                    cell_fill.set(qn('w:fill'), 'E6F0FA')
                    cell._element.get_or_add_tcPr().append(cell_fill)
    else:
        table.rows[1].cells[0].text = "Geen"
    
    for row_idx in range(1, num_rows):
        for cell in table.rows[row_idx].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0, 48, 135)
    
    doc.add_heading("Beskikbare Opvoeders per Periode", level=2)
    for period in teaching_periods:
        period_num = int(period.split()[1].split(" ")[0])
        period_position = teaching_periods.index(period) + 1
        total_periods_first_day = 7 - start_period_idx + 1
        day_to_use = selected_day if period_position <= total_periods_first_day else next_day
        current_absent = [
            t for t in absent_educators
            if st.session_state.return_periods.get(t, "Volle Dag Afwesig") == "Volle Dag Afwesig" or
            (st.session_state.return_periods[t] != "Volle Dag Afwesig" and period_num < int(st.session_state.return_periods[t].split()[-1]))
        ] + [
            t for t in st.session_state.leave_periods.keys()
            if st.session_state.leave_periods[t] != "Geen Vroeë Vertrek" and period_num >= int(st.session_state.leave_periods[t].split()[-1])
        ]
        available = [
            t for t in data[day_to_use][f"Period {period_num}"]
            if t not in current_absent and t not in EXCLUDED_TEACHERS and t not in st.session_state.daily_substitutes[period]
        ]
        doc.add_paragraph(f"{period}: {', '.join(available) if available else 'Geen'}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    st.markdown("<div class='download-button'>", unsafe_allow_html=True)
    st.download_button(
        label="Download TOESIGROOSTER",
        data=buffer,
        file_name="vervangingskedule.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.markdown("</div>", unsafe_allow_html=True)
st.markdown("</div>", unsafe_allow_html=True)
st.markdown("</div>", unsafe_allow_html=True)

# Insights Section
st.markdown("<div class='section'>", unsafe_allow_html=True)
st.header("INSIGTE EN VISUALISERING")

st.subheader("OPVOEDER AFWESIGHEIDS FREKWENSIE (ALGEHEEL)")
absence_data = pd.Series(st.session_state.absence_counts)
if not absence_data.empty:
    fig, ax = plt.subplots(figsize=(12, 6))
    absence_data.plot(kind='bar', ax=ax, color='#005B99')
    plt.title("Aantal Afwesighede per Opvoeder", fontsize=14)
    plt.xlabel("Opvoeder", fontsize=12)
    plt.ylabel("Aantal Afwesighede", fontsize=12)
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    st.pyplot(fig)
else:
    st.write("Geen afwesigheidsdata beskikbaar nie.")

st.subheader("AFWESIGHEIDSOPSOMMING")
summary_period = st.selectbox("Kies Opsommingsperiode", ["Weekliks", "Maandelikse", "Kwartaalliks"], key="summary_period")
current_date = datetime.now()
if summary_period == "Weekliks":
    start_date = current_date - timedelta(days=7)
    period_label = "Laaste 7 Dae"
elif summary_period == "Maandelikse":
    start_date = current_date - timedelta(days=30)
    period_label = "Laaste 30 Dae"
else:
    start_date = current_date - timedelta(days=90)
    period_label = "Laaste 90 Dae"

period_absences = defaultdict(int)
for educator, timestamps in st.session_state.absence_timestamps.items():
    for ts in timestamps:
        if ts >= start_date:
            period_absences[educator] += 1

st.subheader(f"OPVOEDER AFWESIGHEIDS FREKWENSIE ({period_label})")
period_absence_data = pd.Series(period_absences)
if not period_absence_data.empty:
    fig, ax = plt.subplots(figsize=(12, 6))
    period_absence_data.plot(kind='bar', ax=ax, color='#005B99')
    plt.title(f"Aantal Afwesighede per Opvoeder ({period_label})", fontsize=14)
    plt.xlabel("Opvoeder", fontsize=12)
    plt.ylabel("Aantal Afwesighede", fontsize=12)
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    st.pyplot(fig)
else:
    st.write(f"Geen afwesigheidsdata vir die {period_label.lower()} beskikbaar nie.")

st.markdown("<div class='section'>", unsafe_allow_html=True)
st.subheader(f"Opvoeder Frekwensie ({period_label})")
period_usage = defaultdict(int)
for educator, timestamps in st.session_state.usage_timestamps.items():
    for ts, _ in timestamps:
        if ts >= start_date:
            period_usage[educator] += 1

usage_data = pd.Series(period_usage)
if not usage_data.empty:
    usage_df = pd.DataFrame(usage_data.items(), columns=["Opvoeder", "Aantal Vervangings"])
    usage_df = usage_df.sort_values(by="Aantal Vervangings", ascending=False).reset_index(drop=True)
    usage_df.index = usage_df.index + 1
    st.table(usage_df.style.set_properties(**{
        'background-color': '#E6F0FA',
        'color': '#003087',
        'border': '1px solid #005B99',
        'padding': '10px',
        'text-align': 'center',
        'font-size': '16px'
    }).set_table_styles([{
        'selector': 'th',
        'props': [('background-color', '#005B99'), ('color', 'white'), ('font-weight', 'bold'), ('font-size', '16px'), ('padding', '10px')]
    }]))
else:
    st.write(f"Geen vervangingsdata vir die {period_label.lower()} beskikbaar nie.")
st.markdown("</div>", unsafe_allow_html=True)
st.markdown("</div>", unsafe_allow_html=True)